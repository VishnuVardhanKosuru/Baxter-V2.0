"""
agents/doc_parser.py
--------------------
Stage 1 of the Baxter pipeline: turn a folder of Word documents into enriched,
requirement-mapped test case knowledge JSON.

This module is organized by execution flow (how parse_documents() calls it):

  1. Text utilities         Text splitting and normalization (used by parsers)
  2. Document discovery     Scan input_modules/ and pair FRDs with test cases
  3. .docx extraction       Parse Word XML into DocumentAST and TestCaseModel
  4. Compression/summaries  Compact FRD index and test case summaries for LLM
  5. LLM mapping            Map test cases to FRD sections (AI or keyword fallback)
  6. Enrichment             Attach requirement context and auto-generate tags
  7. Module overview        Extract scope, glossary, and project metadata
  8. Orchestration          parse_documents() — main entry point

Execution flow in parse_documents():
  Discover modules -> Parse FRD -> Compress FRD -> Parse test cases ->
  Compress test cases -> LLM mapping (or fallback) -> Enrich -> Write JSON

Token efficiency: compact summaries cut mapping prompt size by ~25-30%. Requests
are load-balanced across all configured API keys via the LiteLLM router.

Public API:
  parse_documents(modules_dir, out_dir, ...) -> List[str]
"""

import asyncio
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import docx
from docx.table import Table as DocxTable
from langchain_core.prompts import ChatPromptTemplate

import core.constants as const
from core.llm_factory import collect_keys, create_llm
from core.logger import logger
from core.models import (
    BatchMappingResponse,
    DocumentAST,
    FeatureContextModel,
    MappedContextModel,
    MappedRef,
    ModuleOverviewModel,
    ParsedDocumentResponse,
    ParserSummaryModel,
    SectionNode,
    TestCaseMapping,
    TestCaseModel,
)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. TEXT UTILITIES
# ═══════════════════════════════════════════════════════════════════════════════

def split_list_field(text: str) -> List[str]:
    """
    Split a delimited field value into a clean list of items.

    Splits on semicolons or newlines — the two delimiter styles used in FRD
    tables for multi-value fields (pre-conditions, business rules).

    Args:
        text: Raw cell text that may contain multiple values, or None.

    Returns:
        List of non-empty, stripped strings with trailing commas removed.
    """
    parts = const.REGEX_DELIMITER_SPLIT.split(text or "")
    return [cleaned for cleaned in (p.strip().strip(",").strip() for p in parts) if cleaned]


def split_numbered_steps(text: str) -> List[str]:
    """
    Split a numbered-step string into an ordered list of step strings.

    Example::

        "1. Open login page 2. Enter credentials" -> ["Open login page", "Enter credentials"]

    Args:
        text: Raw multi-step string from a table cell, or None.

    Returns:
        List of individual step strings with whitespace normalised.
    """
    cleaned = const.REGEX_WHITESPACE.sub(" ", (text or "").strip())
    return [p.strip() for p in const.REGEX_NUMBERED_STEPS.split(cleaned) if p.strip()]


# ═══════════════════════════════════════════════════════════════════════════════
# 2. DOCUMENT DISCOVERY — scan modules and pair FRDs with test cases
# ═══════════════════════════════════════════════════════════════════════════════
# Called first in parse_documents() to find all module folders and documents.
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class ModulePackage:
    """Represents a discovered module folder with classified .docx files."""
    module_folder: str
    frd_files: List[Path] = field(default_factory=list)
    tc_files: List[Path] = field(default_factory=list)


class DocumentClassifier:
    """Classifies .docx files within a module folder as FRD or Test Case suite."""

    @staticmethod
    def classify_files(folder_path: Path) -> Tuple[List[Path], List[Path]]:
        """
        Scan a folder and split .docx files into FRD and TC lists.

        Classification is based on tiered signals (high, medium, low) using
        FRD_FILENAME_SIGNALS and TC_FILENAME_SIGNALS from core.constants,
        to detect conflicts. Word lock/temp files (``~$*``) are skipped.

        Args:
            folder_path: Path to the module sub-folder to scan.

        Returns:
            A 2-tuple (frd_files, tc_files) of sorted Path lists. Both are empty
            if the folder is unreadable.
        """
        frd_files: List[Path] = []
        tc_files: List[Path] = []

        try:
            candidates = sorted(folder_path.glob(f"*{const.SUPPORTED_DOC_EXT}"))
        except OSError as exc:
            logger.error("Cannot read folder %s: %s", folder_path, exc)
            return frd_files, tc_files

        for file in candidates:
            if file.name.startswith(const.WORD_TEMP_PREFIX):
                continue   # Microsoft Word lock/temp file

            name_lower = file.name.lower()

            def get_score(signals: dict, name: str) -> int:
                if any(kw in name for kw in signals["high"]): return 3
                if any(kw in name for kw in signals["medium"]): return 2
                if any(kw in name for kw in signals["low"]): return 1
                return 0

            frd_score = get_score(const.FRD_FILENAME_SIGNALS, name_lower)
            tc_score = get_score(const.TC_FILENAME_SIGNALS, name_lower)

            if frd_score == 0 and tc_score == 0:
                logger.warning("Unclassified document (no keyword match): %s", file)
                continue

            if frd_score > tc_score:
                frd_files.append(file)
            elif tc_score > frd_score:
                tc_files.append(file)
            else:
                logger.warning("Document classification conflict (FRD vs TC score tied at %d): %s. Skipping.", frd_score, file)

        return frd_files, tc_files


class ModuleFolderScanner:
    """Scans the root modules directory and returns all module packages."""

    def __init__(self, root_dir: Path):
        """
        Args:
            root_dir: Path to the top-level input_modules directory.
        """
        self.root_dir = Path(root_dir)

    def scan(self) -> List[ModulePackage]:
        """
        Walk root_dir and return a ModulePackage for every location containing at
        least one classifiable .docx file (the root itself, plus each sub-folder).

        Returns:
            List of ModulePackage objects in deterministic (sorted) order.
            Empty if root_dir does not exist.
        """
        packages: List[ModulePackage] = []
        if not self.root_dir.is_dir():
            logger.error("Modules directory not found: %s", self.root_dir)
            return packages

        # 1. The root directory itself may hold FRD / TC documents directly.
        root_frd, root_tc = DocumentClassifier.classify_files(self.root_dir)
        if root_frd or root_tc:
            packages.append(ModulePackage(
                module_folder=self.root_dir.name or "input_modules",
                frd_files=root_frd,
                tc_files=root_tc,
            ))

        # 2. Each sub-folder is its own module. sorted() keeps processing order
        #    deterministic — filesystem iteration order is not guaranteed.
        try:
            entries = sorted(self.root_dir.iterdir())
        except OSError as exc:
            logger.error("Cannot list %s: %s", self.root_dir, exc)
            return packages

        for item in entries:
            if not item.is_dir():
                continue
            frd_files, tc_files = DocumentClassifier.classify_files(item)
            if frd_files or tc_files:
                packages.append(ModulePackage(
                    module_folder=item.name,
                    frd_files=frd_files,
                    tc_files=tc_files,
                ))

        return packages


# ═══════════════════════════════════════════════════════════════════════════════
# 3. EXTRACTION — parse FRD and test case documents into structured data
# ═══════════════════════════════════════════════════════════════════════════════
# FRDModuleParser.parse() is called second; TestCaseModuleParser.parse() third.
# Both walk the Word XML, recognize patterns (headings, tables, keywords), and
# extract into DocumentAST (FRD tree) and TestCaseModel objects (test list).
# ═══════════════════════════════════════════════════════════════════════════════

class FRDModuleParser:
    """Extracts a .docx FRD into a structured DocumentAST."""

    _PSTYLE_XPATH = f"{const.XML_WORD_NAMESPACE}pPr/{const.XML_WORD_NAMESPACE}pStyle"
    _STYLE_VAL_ATTR = f"{const.XML_WORD_NAMESPACE}val"

    @staticmethod
    def _derive_section_id(title: str, index: int) -> Tuple[str, str]:
        """
        Derive a (section_id_suffix, section_type) pair from a heading title.

        Requirement headings are matched first and keep their FR-nnn number when
        present. Everything else is matched against SECTION_TYPE_RULES in order,
        falling back to a slug of the heading text.

        Args:
            title: Raw heading text from the paragraph.
            index: Number of sections already parsed (used as a fallback ID).

        Returns:
            (id_suffix, type) where type is one of: functional, nfr, interface,
            scope, glossary, general.
        """
        title_lower = title.lower()

        if const.HEADING_REQUIREMENT_KEYWORD.lower() in title_lower:
            match = const.REGEX_REQUIREMENT_ID.search(title)
            if match:
                return match.group(1).strip(), "functional"
            return f"FR-{str(index).zfill(3)}", "functional"

        for keywords, suffix, sec_type in const.SECTION_TYPE_RULES:
            if any(kw in title_lower for kw in keywords):
                return suffix, sec_type

        clean_title = const.REGEX_LEADING_NUMBER.sub("", title).strip()
        slug = const.REGEX_NON_WORD.sub("_", clean_title.lower()).strip("_")
        slug = slug[:const.SECTION_SLUG_MAX_LEN].upper().strip("_")
        return slug or f"SEC-{index}", "general"

    @staticmethod
    def _extract_table_metadata(table: DocxTable, meta: dict) -> List[List[str]]:
        """
        Populate `meta` from a key/value FRD table and return its raw rows.

        The first cell of each row is treated as the field label and the second
        as its value; unrecognised labels are kept in the returned raw rows only
        (used later for glossary extraction).
        """
        table_rows: List[List[str]] = []

        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2:
                continue

            key = cells[0].lower().rstrip(":")
            val = cells[1]

            if const.KEY_DESCRIPTION in key:
                meta["description"] = val
            elif const.KEY_ACTORS in key:
                meta["actors"] = [a.strip() for a in const.REGEX_TYPE_SPLIT.split(val) if a.strip()]
            elif const.KEY_TRIGGER in key:
                meta["trigger"] = val
            elif const.KEY_PRIORITY in key:
                meta["priority"] = val
            elif any(k in key for k in const.KEY_PRECONDITIONS):
                meta["pre_conditions"] = split_list_field(val)
            elif const.KEY_BUSINESS_RULES in key:
                meta["business_rules"] = split_list_field(val)
            elif any(k in key for k in const.KEY_EXCEPTION_FLOW):
                meta["exception_flows"] = split_numbered_steps(val)
            elif const.KEY_MAIN_FLOW in key:
                meta["main_flow"] = split_numbered_steps(val)
            elif any(k in key for k in const.KEY_POSTCONDITIONS):
                meta["post_conditions"] = split_list_field(val)

            table_rows.append(cells)

        return table_rows

    @classmethod
    def parse(cls, file_path: Path, module_folder: str) -> DocumentAST:
        """
        Parse a single FRD .docx file into a DocumentAST.

        The parser walks the document body element-by-element:
          - Paragraph with a recognised heading -> start a new SectionNode.
          - Table following a section heading   -> populate that section's metadata.
          - All other paragraphs                -> appended to the current section.

        Section IDs are prefixed with the module folder (e.g.
        "01_User_Auth:FR-001") and de-duplicated with a numeric suffix, so two
        headings deriving the same ID cannot collide and overwrite each other in
        downstream section lookups.

        Args:
            file_path:     Path to the FRD .docx file.
            module_folder: Name of the parent module folder (used in section IDs).

        Returns:
            A DocumentAST with all parsed sections. Contains only the seeded
            header section if the file cannot be opened.
        """
        file_path = Path(file_path)
        ast = DocumentAST(module_folder=module_folder, source_file=file_path.name)

        try:
            doc = docx.Document(str(file_path))
        except Exception as exc:
            # python-docx raises a variety of types (PackageNotFoundError, KeyError,
            # zipfile.BadZipFile) for a malformed or non-.docx file. A bad input
            # document must degrade to an empty AST, never crash the pipeline.
            logger.error("Failed to read FRD docx %s: %s", file_path.name, exc)
            return ast

        # Seed a generic header section to catch paragraphs before the first heading.
        current_section = SectionNode(
            section_id=f"{module_folder}:GEN-001",
            title="Document Header & Overview",
            type="general",
            module_folder=module_folder,
            source_file=file_path.name,
        )
        ast.sections.append(current_section)
        seen_ids = {current_section.section_id}

        for child in doc.element.body:
            local_tag = child.tag.split("}")[-1]

            if local_tag == const.XML_PARAGRAPH_TAG:
                p_text = "".join(
                    node.text or ""
                    for node in child.iter()
                    if node.tag.endswith(const.XML_TEXT_TAG_SUFFIX)
                ).strip()
                if not p_text:
                    continue

                # A heading is signalled either by a Word Heading* style or by
                # the requirement-ID keyword appearing in the text.
                is_heading = False
                style_node = child.find(cls._PSTYLE_XPATH)
                if style_node is not None and "Heading" in style_node.attrib.get(cls._STYLE_VAL_ATTR, ""):
                    is_heading = True

                if const.HEADING_REQUIREMENT_KEYWORD in p_text or is_heading:
                    sec_id_suffix, sec_type = cls._derive_section_id(p_text, len(ast.sections))
                    sec_id = f"{module_folder}:{sec_id_suffix}"

                    # De-duplicate: "Scope" and "Out of Scope" both derive SCOPE,
                    # which would otherwise silently overwrite one another.
                    if sec_id in seen_ids:
                        dedupe = 2
                        while f"{sec_id}-{dedupe}" in seen_ids:
                            dedupe += 1
                        sec_id = f"{sec_id}-{dedupe}"
                    seen_ids.add(sec_id)

                    current_section = SectionNode(
                        section_id=sec_id,
                        title=p_text,
                        type=sec_type,
                        module_folder=module_folder,
                        source_file=file_path.name,
                    )
                    ast.sections.append(current_section)
                else:
                    current_section.paragraphs.append(p_text)

            elif local_tag == const.XML_TABLE_TAG:
                table_rows = cls._extract_table_metadata(
                    DocxTable(child, doc), current_section.metadata
                )
                if table_rows:
                    current_section.tables.append(table_rows)

        return ast


class TestCaseModuleParser:
    """Extracts all test cases from all tables in a .docx file."""

    @staticmethod
    def _map_columns(headers: List[str]) -> Dict[str, Optional[int]]:
        """
        Resolve logical column names to physical column indices.

        A target shorter than COL_SUBSTRING_MIN_LEN must match a header exactly;
        longer targets may match as a substring. This stops short keys such as
        "tc" matching unrelated headers like "structure".
        """
        def find_col(targets: Tuple[str, ...]) -> Optional[int]:
            for idx, header in enumerate(headers):
                for target in targets:
                    if target == header or (
                        len(target) > const.COL_SUBSTRING_MIN_LEN and target in header
                    ):
                        return idx
            return None

        return {
            "test_name":        find_col(const.COL_TEST_NAME),
            "type":             find_col(const.COL_TYPE),
            "subject":          find_col(const.COL_SUBJECT),
            "description":      find_col(const.COL_DESCRIPTION),
            "expected_result":  find_col(const.COL_EXPECTED_RESULT),
            "execution_status": find_col(const.COL_EXECUTION_STATUS),
        }

    @staticmethod
    def _cell(cells: List[str], col_map: Dict[str, Optional[int]], key: str) -> str:
        """Safe column accessor — returns "" for missing or out-of-range columns."""
        idx = col_map.get(key)
        if idx is None or idx >= len(cells):
            return ""
        return cells[idx]

    @classmethod
    def parse(cls, file_path: Path, module_folder: str) -> List[TestCaseModel]:
        """
        Parse a Manual Test Cases .docx file into TestCaseModel objects.

        Every table is inspected: row 0 is the header row, later rows are data.
        Tables without a recognisable "test name" column are skipped. Column
        discovery is dynamic and supports the alternative header spellings
        defined in core.constants.

        Args:
            file_path:     Path to the .docx Test Cases file.
            module_folder: Name of the parent module folder (stored on each TC).

        Returns:
            Ordered list of TestCaseModel instances. Empty if the file cannot be
            opened or contains no valid test case tables.
        """
        file_path = Path(file_path)
        test_cases: List[TestCaseModel] = []

        try:
            doc = docx.Document(str(file_path))
        except Exception as exc:
            # See FRDModuleParser.parse — python-docx raises varied types for
            # malformed input; degrade to an empty list instead of crashing.
            logger.error("Failed to read TC docx %s: %s", file_path.name, exc)
            return test_cases

        if not doc.tables:
            logger.warning("No tables found in test case document %s.", file_path.name)
            return test_cases

        for table_idx, table in enumerate(doc.tables):
            if not table.rows:
                continue

            headers = [(cell.text or "").strip().lower() for cell in table.rows[0].cells]
            col_map = cls._map_columns(headers)

            if col_map["test_name"] is None:
                logger.debug(
                    "Table %d in %s has no 'test name' column — skipped.", table_idx, file_path.name
                )
                continue

            for row_idx, row in enumerate(table.rows[1:], start=1):
                cells = [(cell.text or "").strip() for cell in row.cells]

                tc_name_full = cls._cell(cells, col_map, "test_name")
                if not tc_name_full:
                    continue

                tc_id_match = const.REGEX_TC_ID.search(tc_name_full)
                tc_id = tc_id_match.group(1).upper() if tc_id_match else f"TC-UNKN-{table_idx}-{row_idx}"

                title_match = const.REGEX_TC_TITLE.search(tc_name_full)
                title = title_match.group(1).strip() if title_match else tc_name_full

                type_str = cls._cell(cells, col_map, "type")
                types = [t.strip() for t in const.REGEX_TYPE_SPLIT.split(type_str) if t.strip()]

                desc_text = cls._cell(cells, col_map, "description")
                steps = split_numbered_steps(desc_text)
                if not steps and desc_text:
                    steps = [desc_text]

                test_cases.append(TestCaseModel(
                    tc_id=tc_id,
                    title=title,
                    module_folder=module_folder,
                    source_file=file_path.name,
                    source_table_index=table_idx,
                    source_row_index=row_idx,
                    type=types,
                    subject=cls._cell(cells, col_map, "subject"),
                    execution_status=cls._cell(cells, col_map, "execution_status"),
                    steps=steps,
                    expected_result=cls._cell(cells, col_map, "expected_result"),
                ))

        return test_cases


# ═══════════════════════════════════════════════════════════════════════════════
# 4. COMPRESSION & SUMMARY — compact FRD and test cases for LLM efficiency
# ═══════════════════════════════════════════════════════════════════════════════
# After extraction, build_compact_section_index() creates a one-line index of
# each FRD section (~25-30% smaller than raw). Test case summaries are built
# inline in parse_documents(). These summaries are passed to the LLM.
#
# Then: LLM MAPPING
# ─────────────────
# _map_all_modules() sends summaries to the LLM in one batched request per module.
# The LLM answers: "Which FRD section does each test case exercise?"
# If LLM unavailable or fails, _generate_fallback_mappings() uses keyword matching.
# ═══════════════════════════════════════════════════════════════════════════════

MAPPER_SYSTEM_PROMPT = (
    "You are an expert Test Automation Architect. "
    "Map each Manual Test Case to ALL relevant FRD Section IDs from the provided Index. "
    "A test case may map to multiple sections (e.g. Functional Requirement, NFR, Scope). "
    "Return all mapped references per test case with confidence score (0.0-1.0) and a 1-sentence reason."
)

MAPPER_PROMPT = ChatPromptTemplate.from_messages([
    ("system", MAPPER_SYSTEM_PROMPT),
    ("user", "FRD Index:\n{index}\n\nTest Cases to Map:\n{test_cases}")
])

_mapper_chain = None


def reset_mapper_chain() -> None:
    """
    Discards the cached mapper chain so the next call rebuilds it.

    Called after the set of available API keys changes at runtime (see
    server._inject_ui_gemini_key) — without this the chain would keep routing
    through the key set captured at first use.
    """
    global _mapper_chain
    _mapper_chain = None


def _get_mapper_chain():
    """
    Lazily builds the multi-key LiteLLM mapper chain.

    Returns None when no usable LLM is configured, which signals callers to fall
    back to heuristic mapping rather than failing the whole parse.
    """
    global _mapper_chain
    if _mapper_chain is None:
        try:
            bundle = create_llm()
            _mapper_chain = MAPPER_PROMPT | bundle.llm.with_structured_output(BatchMappingResponse)
        except (RuntimeError, ImportError, ValueError) as exc:
            logger.warning("LLM mapper unavailable (%s). Heuristic fallback mapping will be used.", exc)
            return None
    return _mapper_chain


def build_compact_section_index(ast: DocumentAST) -> str:
    """
    Builds a token-efficient compact index of the FRD for LLM mapping.

    One line per section: `[section_id] title (type) | first-paragraph summary`.
    Total output is capped at COMPACT_INDEX_MAX_CHARS so a very large FRD cannot
    blow up the prompt (and the cost of every mapping call that uses it).
    """
    lines = [f"Module: {ast.module_folder}"]
    budget = const.COMPACT_INDEX_MAX_CHARS
    used = len(lines[0])
    truncated = 0

    for section in ast.sections:
        line = f"[{section.section_id}] {section.title} ({section.type})"
        if section.paragraphs:
            first_p = section.paragraphs[0].strip()
            if first_p:
                summary = first_p[:const.COMPACT_SUMMARY_CHARS]
                if len(first_p) > const.COMPACT_SUMMARY_CHARS:
                    summary += "..."
                line += f" | {summary}"

        if used + len(line) + 1 > budget:
            truncated += 1
            continue
        lines.append(line)
        used += len(line) + 1

    if truncated:
        logger.warning(
            "Compact index for %s hit the %d char budget — %d section(s) omitted.",
            ast.module_folder, budget, truncated,
        )

    return "\n".join(lines)


def _generate_fallback_mappings(module_tcs: List[TestCaseModel], ast: DocumentAST) -> BatchMappingResponse:
    """
    Heuristic keyword mapper used when the LLM is unavailable or rate limited.

    Scores each FRD section by how many significant words of its title appear in
    the test case text, then keeps only the best-scoring sections so a test case
    is not indiscriminately linked to every section in the document.
    """
    mappings = []
    section_ids = [s.section_id for s in ast.sections if s.section_id]

    # Pre-compute significant title words once per section instead of per TC.
    section_words = [
        (sec, {w for w in sec.title.lower().split() if len(w) > const.FALLBACK_MIN_WORD_LEN})
        for sec in ast.sections
    ]

    for tc in module_tcs:
        tc_text = f"{tc.title} {tc.subject} {' '.join(tc.steps)}".lower()

        scored = []
        for sec, words in section_words:
            if not words:
                continue
            hits = sum(1 for w in words if w in tc_text)
            if hits:
                scored.append((hits, sec))

        matched_refs = []
        if scored:
            best_score = max(hits for hits, _ in scored)
            for hits, sec in scored:
                if hits == best_score:
                    matched_refs.append(MappedRef(
                        ref_id=sec.section_id,
                        confidence=const.FALLBACK_CONFIDENCE,
                        reason=f"Keyword overlap with FRD section: {sec.title}",
                    ))

        if not matched_refs and section_ids:
            matched_refs.append(MappedRef(
                ref_id=section_ids[0],
                confidence=const.FALLBACK_DEFAULT_CONFIDENCE,
                reason="Default requirement mapping — no keyword overlap found",
            ))

        mappings.append(TestCaseMapping(tc_id=tc.tc_id, mapped_refs=matched_refs))

    return BatchMappingResponse(mappings=mappings)


def _compute_mapping_concurrency() -> int:
    """Concurrency for the mapping batch: RPM x active keys, capped for memory."""
    rpm = const.env_int("GEMINI_RPM", 50)
    num_keys = len(collect_keys("GEMINI_API_KEY"))
    if num_keys <= 0:
        return rpm
    return min(rpm * num_keys, const.LLM_MAX_CONCURRENCY)


def _map_all_modules(batch_inputs: List[dict]) -> list:
    """
    Runs the LLM mapping call for every module concurrently.

    Returns a list positionally aligned with batch_inputs. Failures are returned
    as Exception objects (never raised) so one bad module cannot abort the run —
    callers substitute heuristic mappings for those entries.
    """
    chain = _get_mapper_chain()
    if chain is None:
        return [RuntimeError("No LLM configured")] * len(batch_inputs)

    concurrency = _compute_mapping_concurrency()
    logger.info(
        "[MAP] Batch mapping %d module(s) concurrently (max_concurrency=%d)...",
        len(batch_inputs), concurrency,
    )

    async def _run() -> list:
        return await chain.abatch(
            batch_inputs,
            config={"max_concurrency": concurrency},
            return_exceptions=True,
        )

    try:
        return asyncio.run(_run())
    except Exception as exc:
        logger.error("Batch LLM mapping failed: %s. Using heuristic fallback mapping.", exc)
        return [exc] * len(batch_inputs)


# ═══════════════════════════════════════════════════════════════════════════════
# 5. ENRICHMENT — attach requirement context and auto-generate tags
# ═══════════════════════════════════════════════════════════════════════════════
# After mapping, enrich_module_test_cases() takes each test case and:
#   1. Finds its best-matched FRD section (from the LLM or fallback)
#   2. Copies that section's full context (description, actors, business rules...)
#   3. Builds automatic tags (@Login, @Positive, @REQ-42, etc.)
#
# Also here: build_module_overview() extracts project scope, glossary, metadata.
# ═══════════════════════════════════════════════════════════════════════════════

def _build_feature_context(sec: SectionNode) -> FeatureContextModel:
    """Projects a SectionNode's metadata into the LLM-facing feature context."""
    meta = sec.metadata or {}
    description = meta.get("description", "")
    if not description and sec.paragraphs:
        description = "\n".join(sec.paragraphs[:2])

    return FeatureContextModel(
        feature_name=sec.title,
        description=description,
        trigger=meta.get("trigger", ""),
        priority=meta.get("priority", ""),
        actors=meta.get("actors", []),
        pre_conditions=meta.get("pre_conditions", []),
        main_flow=meta.get("main_flow", []),
        post_conditions=meta.get("post_conditions", []),
        business_rules=meta.get("business_rules", []),
        exception_flows=meta.get("exception_flows", []),
    )


def _build_cucumber_tags(tc: TestCaseModel, best_ref_id: str) -> List[str]:
    """Derives a de-duplicated @tag list from the TC type, subject, and FRD ref."""
    tags = [f"@{t.lower()}" for t in tc.type] if tc.type else []

    subject_slug = "".join(c if c.isalnum() else "_" for c in tc.subject.lower()).strip("_")
    if subject_slug:
        tags.append(f"@{subject_slug}")

    tags.append(f"@{best_ref_id.lower().replace('-', '_').replace(':', '_')}")
    return list(dict.fromkeys(tags))


def enrich_module_test_cases(
    test_cases: List[TestCaseModel],
    mapping_response: BatchMappingResponse,
    ast: DocumentAST,
) -> List[TestCaseModel]:
    """
    Merges FRD section context into each test case based on the mapping decision.

    Sets feature_ref to the highest-confidence section, feature_refs to every
    mapped section, mapped_contexts to the full extracted context per section,
    and auto-generates cucumber tags. Test cases with no mapping are marked
    UNKNOWN rather than dropped.
    """
    section_map = {sec.section_id: sec for sec in ast.sections}
    tc_mapping_dict = {m.tc_id: m for m in mapping_response.mappings}

    enriched_tcs = []
    for tc in test_cases:
        mapping = tc_mapping_dict.get(tc.tc_id)
        if not mapping or not mapping.mapped_refs:
            tc.feature_ref = const.UNKNOWN_FEATURE_REF
            enriched_tcs.append(tc)
            continue

        mapped_contexts = []
        feature_refs = []

        for ref in mapping.mapped_refs:
            # The LLM sees section IDs wrapped in brackets in the index, and
            # sometimes echoes them back that way. Normalise before lookup so
            # the ID stored on the test case always matches the AST.
            clean_ref_id = ref.ref_id.strip().strip("[]")
            feature_refs.append(clean_ref_id)
            sec = section_map.get(clean_ref_id)

            mapped_contexts.append(MappedContextModel(
                ref_id=clean_ref_id,
                title=sec.title if sec else "Unknown Section",
                type=sec.type if sec else "general",
                confidence=ref.confidence,
                reason=ref.reason,
                context=_build_feature_context(sec) if sec else None,
            ))

        best_ref = max(mapping.mapped_refs, key=lambda x: x.confidence)
        clean_best_ref_id = best_ref.ref_id.strip().strip("[]")

        tc.feature_ref = clean_best_ref_id
        tc.feature_refs = list(dict.fromkeys(feature_refs))
        tc.mapped_contexts = mapped_contexts
        tc.cucumber_tags = _build_cucumber_tags(tc, clean_best_ref_id)

        enriched_tcs.append(tc)

    return enriched_tcs


def build_module_overview(ast: DocumentAST) -> ModuleOverviewModel:
    """
    Extracts document-level metadata (purpose, scope, glossary) from the AST.

    Scope paragraphs are split into in-scope and out-of-scope based on both the
    section heading and per-paragraph wording.
    """
    overview = ModuleOverviewModel()

    for sec in ast.sections:
        title_lower = sec.title.lower()

        if "purpose" in title_lower and not overview.purpose:
            overview.purpose = "\n".join(sec.paragraphs)

        elif "scope" in title_lower:
            heading_is_out = "out of scope" in title_lower or "out-of-scope" in title_lower
            for p in sec.paragraphs:
                p_lower = p.lower()
                if heading_is_out or "out of scope" in p_lower or "out-of-scope" in p_lower:
                    overview.out_of_scope.append(p)
                else:
                    overview.in_scope.append(p)

        elif "glossary" in title_lower or sec.type == "glossary":
            for table in sec.tables:
                for row in table:
                    if len(row) >= 2:
                        term, definition = row[0].strip(), row[1].strip()
                        if term and term.lower() not in const.GLOSSARY_SKIP_TERMS:
                            overview.glossary[term] = definition

    return overview


# ═══════════════════════════════════════════════════════════════════════════════
# 6. ORCHESTRATION — parse_documents() ties everything together
# ═══════════════════════════════════════════════════════════════════════════════
# Public entry point. Calls section 2 (discover) → section 3 (extract) →
# section 4 (map) → section 5 (enrich) → writes JSON.
# Also handles two calling conventions: directory mode and single-pair mode.
# ═══════════════════════════════════════════════════════════════════════════════

def _resolve_single_pair(frd_path: str, tc_path: str) -> ModulePackage:
    """
    Validates and wraps an explicit FRD + TC file pair as a single ModulePackage.

    Raises:
        FileNotFoundError: if either document does not exist.
    """
    frd_p, tc_p = Path(frd_path), Path(tc_path)

    missing = [str(p) for p in (frd_p, tc_p) if not p.is_file()]
    if missing:
        raise FileNotFoundError(f"Input document(s) not found: {', '.join(missing)}")

    module_name = frd_p.stem.replace("_FRD", "").replace("FRD_", "") or "Module"
    return ModulePackage(module_folder=module_name, frd_files=[frd_p], tc_files=[tc_p])


def parse_documents(
    modules_dir_or_frd: str,
    out_dir_or_tc: str,
    target_out_dir: Optional[str] = None,
    project: str = const.DEFAULT_PROJECT_NAME,
    version: str = const.DEFAULT_VERSION,
    skip_types: Optional[List[str]] = None,
) -> List[str]:
    """
    Orchestrates the parse + map + enrich pipeline.

    Two calling conventions are supported:
      1. Directory mode:   parse_documents("input_modules", "output")
      2. Single-pair mode:  parse_documents("FRD.docx", "TC.docx", "output")
         (selected automatically when the first two arguments are .docx paths)

    Returns:
        List of generated knowledge JSON paths, one per module. Empty if no
        module yielded parsable content.

    Raises:
        FileNotFoundError: in single-pair mode, if either document is missing.
    """
    skip_types = list(skip_types) if skip_types is not None else list(const.DEFAULT_SKIP_TYPES)

    is_single_pair = (
        str(modules_dir_or_frd).lower().endswith(const.SUPPORTED_DOC_EXT)
        and str(out_dir_or_tc).lower().endswith(const.SUPPORTED_DOC_EXT)
    )

    if is_single_pair:
        out_dir = target_out_dir or str(const.DIR_OUTPUT)
        packages = [_resolve_single_pair(modules_dir_or_frd, out_dir_or_tc)]
    else:
        out_dir = out_dir_or_tc
        logger.info("Scanning modules directory: %s", modules_dir_or_frd)
        packages = ModuleFolderScanner(Path(modules_dir_or_frd)).scan()

    if not packages:
        logger.warning("No module packages found in %s.", modules_dir_or_frd)
        return []

    # ── 1. Parse every module and collect LLM mapping inputs ──────────────────
    batch_inputs: List[dict] = []
    packages_to_process = []

    for package in packages:
        logger.info("[MODULE] %s", package.module_folder)

        if not package.frd_files or not package.tc_files:
            logger.warning(
                "[MODULE] %s: missing %s document — skipping.",
                package.module_folder,
                "FRD" if not package.frd_files else "Test Case",
            )
            continue

        frd_file = package.frd_files[0]
        logger.info("  [FRD] Extracting %s", frd_file.name)
        ast = FRDModuleParser.parse(frd_file, package.module_folder)

        compact_index = build_compact_section_index(ast)
        logger.info("  [FRD] Compact index: %d chars, %d sections", len(compact_index), len(ast.sections))

        module_tcs: List[TestCaseModel] = []
        for tc_file in package.tc_files:
            logger.info("  [TC] Extracting from %s", tc_file.name)
            module_tcs.extend(TestCaseModuleParser.parse(tc_file, package.module_folder))

        if skip_types:
            skip_set = {s.lower() for s in skip_types}
            filtered_tcs = [
                tc for tc in module_tcs
                if not any(t.lower() in skip_set for t in tc.type)
            ]
            if len(filtered_tcs) < len(module_tcs):
                logger.info(
                    "  [TC Filter] Filtered out %d non-UI test case(s) matching skip_types %s",
                    len(module_tcs) - len(filtered_tcs),
                    skip_types,
                )
            module_tcs = filtered_tcs

        if not module_tcs:
            logger.warning("[MODULE] %s: no test cases found after filtering — skipping.", package.module_folder)
            continue

        tc_summaries = []
        for tc in module_tcs:
            steps_preview = " ".join(tc.steps)[:const.COMPACT_STEPS_CHARS] if tc.steps else "None"
            types_str = ",".join(tc.type) if tc.type else "General"
            tc_summaries.append(
                f"[{tc.tc_id}] {tc.title} | Subj: {tc.subject} | Type: {types_str} | Steps: {steps_preview}"
            )

        batch_inputs.append({"index": compact_index, "test_cases": "\n".join(tc_summaries)})
        packages_to_process.append((package, ast, module_tcs))

    if not batch_inputs:
        logger.warning("No parsable module content found — nothing to map.")
        return []

    # ── 2. Map test cases to FRD sections (one LLM call per module) ───────────
    responses = _map_all_modules(batch_inputs)

    # ── 3. Enrich and persist one knowledge JSON per module ───────────────────
    knowledge_dir = Path(out_dir) / "knowledge"
    try:
        knowledge_dir.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        raise RuntimeError(f"Cannot create knowledge output directory {knowledge_dir}: {exc}") from exc

    generated_files: List[str] = []

    for idx, (package, ast, module_tcs) in enumerate(packages_to_process):
        response = responses[idx] if idx < len(responses) else RuntimeError("Missing mapping response")

        if isinstance(response, BaseException) or not getattr(response, "mappings", None):
            logger.info("[FALLBACK] Heuristic requirement mapping for %s.", package.module_folder)
            mapping_response = _generate_fallback_mappings(module_tcs, ast)
        else:
            mapping_response = response

        logger.info("[ENRICH] Merging FRD context for %s...", package.module_folder)
        enriched = enrich_module_test_cases(module_tcs, mapping_response, ast)
        overview = build_module_overview(ast)

        module_slug = "".join(
            c if c.isalnum() else "_" for c in package.module_folder.lower()
        ).strip("_") or "module"
        module_out_path = knowledge_dir / f"{module_slug}_knowledge.json"

        module_payload = ParsedDocumentResponse(
            project=f"{project} - {package.module_folder}",
            version=version,
            module_overview=overview,
            summary=ParserSummaryModel(
                total_test_cases=len(enriched),
                skipped_types=skip_types,
            ),
            test_cases=enriched,
        )

        try:
            with open(module_out_path, "w", encoding="utf-8") as f:
                json.dump(module_payload.to_dict(), f, indent=2, ensure_ascii=False)
            logger.info("[SAVED] %s (%d test cases)", module_out_path, len(enriched))
            generated_files.append(str(module_out_path))
        except (OSError, TypeError, ValueError) as exc:
            logger.error("Failed to save module JSON %s: %s", module_out_path, exc)

    logger.info("[SUCCESS] Parsing complete — %d module JSON file(s) in %s",
                len(generated_files), knowledge_dir)

    return generated_files
