"""
agents/scanners.py
------------------
Implements folder scanning, document classification, and multi-file .docx
extraction for the Baxter Version 2 Parser Engine.

Classes
-------
  ModulePackage          -- dataclass: one discovered module folder
  DocumentClassifier     -- classifies .docx files as FRD or TC
  ModuleFolderScanner    -- walks the input_modules directory
  FRDModuleParser        -- extracts a .docx FRD into a DocumentAST
  TestCaseModuleParser   -- extracts test cases from all tables in a .docx

Functions
---------
  split_list_field()     -- splits semicolon/newline delimited cell values
  split_numbered_steps() -- splits numbered step strings into ordered lists
"""

import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from dataclasses import dataclass

import docx
from docx.table import Table as DocxTable

from core import constants as const
from core.models import (
    SectionNode,
    DocumentAST,
    TestCaseModel,
)


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class ModulePackage:
    """Represents a discovered module folder with classified .docx files."""
    module_folder: str
    frd_files: List[Path]
    tc_files: List[Path]


# ---------------------------------------------------------------------------
# Document classification
# ---------------------------------------------------------------------------

class DocumentClassifier:
    """Classifies .docx files within a module folder as FRD or Test Case suite."""

    @staticmethod
    def classify_files(folder_path: Path) -> Tuple[List[Path], List[Path]]:
        """
        Scan a folder and split .docx files into FRD and TC lists.

        Classification is keyword-based (case-insensitive, partial-match)
        using the FRD_FILENAME_KEYWORDS and TC_FILENAME_KEYWORDS lists from
        core.constants.  Files whose names start with "~$" (Word temp files)
        are silently skipped.

        Args:
            folder_path: Path to the module sub-folder to scan.

        Returns:
            A 2-tuple (frd_files, tc_files) where each element is a list of
            Path objects for the matching .docx files.
        """
        frd_files: List[Path] = []
        tc_files:  List[Path] = []

        for file in folder_path.glob("*.docx"):
            if file.name.startswith("~$"):
                continue   # Skip Microsoft Word lock/temp files

            name_lower = file.name.lower()
            if any(kw in name_lower for kw in const.FRD_FILENAME_KEYWORDS):
                frd_files.append(file)
            elif any(kw in name_lower for kw in const.TC_FILENAME_KEYWORDS):
                tc_files.append(file)

        return frd_files, tc_files


# ---------------------------------------------------------------------------
# Module folder scanner
# ---------------------------------------------------------------------------

class ModuleFolderScanner:
    """Scans the root modules directory and returns all module packages."""

    def __init__(self, root_dir: Path):
        """
        Args:
            root_dir: Path to the top-level input_modules directory.
        """
        self.root_dir = Path(root_dir)

    def scan(self) -> List[ModulePackage]:
        """
        Walk root_dir and return a ModulePackage for every sub-folder that
        contains at least one classifiable .docx file.

        Returns:
            Ordered list of ModulePackage objects (order = filesystem order).
        """
        # 1. Check if root directory itself contains FRD / TC documents
        root_frd, root_tc = DocumentClassifier.classify_files(self.root_dir)
        if root_frd or root_tc:
            packages.append(ModulePackage(
                module_folder=self.root_dir.name or "input_modules",
                frd_files=root_frd,
                tc_files=root_tc,
            ))

        # 2. Check each sub-folder
        for item in sorted(self.root_dir.iterdir()):   # sorted for determinism
            if item.is_dir():
                frd_files, tc_files = DocumentClassifier.classify_files(item)
                if frd_files or tc_files:
                    packages.append(ModulePackage(
                        module_folder=item.name,
                        frd_files=frd_files,
                        tc_files=tc_files,
                    ))
        return packages


# ---------------------------------------------------------------------------
# Text utilities
# ---------------------------------------------------------------------------

def split_list_field(text: str) -> List[str]:
    """
    Split a delimited field value into a clean list of items.

    Splits on semicolons or newline characters -- the two delimiter styles
    used in FRD tables for multi-value fields (pre-conditions, business rules).

    Args:
        text: Raw cell text that may contain multiple values.

    Returns:
        List of non-empty, stripped strings with trailing commas removed.
    """
    parts = const.REGEX_DELIMITER_SPLIT.split(text or "")
    return [p.strip().strip(",").strip() for p in parts if p.strip().strip(",").strip()]


def split_numbered_steps(text: str) -> List[str]:
    """
    Split a numbered-step string into an ordered list of step strings.

    Example::

        "1. Open login page 2. Enter credentials" -> ["Open login page", "Enter credentials"]

    Args:
        text: Raw multi-step string from a table cell (may contain newlines).

    Returns:
        List of individual step strings with whitespace normalised.
    """
    cleaned = const.REGEX_WHITESPACE.sub(" ", (text or "").strip())
    return [p.strip() for p in const.REGEX_NUMBERED_STEPS.split(cleaned) if p.strip()]


# ---------------------------------------------------------------------------
# FRD parser
# ---------------------------------------------------------------------------

class FRDModuleParser:
    """Extracts a .docx FRD into a structured DocumentAST."""

    @staticmethod
    def parse(file_path: Path, module_folder: str) -> DocumentAST:
        """
        Parse a single FRD .docx file into a DocumentAST.

        The parser walks the document body element-by-element:
          - Paragraph with a recognised heading -> start a new SectionNode.
          - Table following a section heading  -> populate the section metadata.
          - All other paragraphs               -> appended to current section.

        Section IDs are derived from the heading text using ``derive_section_id``
        and prefixed with the module_folder (e.g. "01_User_Auth:FR-001").

        Args:
            file_path:     Path to the FRD .docx file.
            module_folder: Name of the parent module folder (used in section IDs).

        Returns:
            A DocumentAST containing all parsed SectionNode objects.
        """
        ast = DocumentAST(module_folder=module_folder, source_file=file_path.name)
        try:
            doc = docx.Document(str(file_path))
        except Exception as exc:
            print(f"  [ERROR] Failed to read FRD docx {file_path.name}: {exc}")
            return ast

        # Seed with a generic header section to catch pre-requirement paragraphs.
        current_section = SectionNode(
            section_id=f"{module_folder}:GEN-001",
            title="Document Header & Overview",
            type="general",
            module_folder=module_folder,
            source_file=file_path.name,
        )
        ast.sections.append(current_section)

        def derive_section_id(title: str, index: int) -> Tuple[str, str]:
            """
            Derive a (section_id_suffix, section_type) pair from a heading title.

            Args:
                title: Raw heading text from the paragraph.
                index: Current number of sections already parsed (used as fallback ID).

            Returns:
                A 2-tuple: (id_suffix: str, type: str).
                type is one of: "functional", "nfr", "interface", "scope", "glossary", "general".
            """
            title_lower = title.lower()
            if const.HEADING_REQUIREMENT_KEYWORD.lower() in title_lower:
                match = const.REGEX_REQUIREMENT_ID.search(title)
                if match:
                    return match.group(1).strip(), "functional"
                return f"FR-{str(index).zfill(3)}", "functional"
            elif "scope" in title_lower:
                return "SCOPE", "scope"
            elif "purpose" in title_lower:
                return "PURPOSE", "general"
            elif "interface" in title_lower:
                return "INTF", "interface"
            elif "performance" in title_lower:
                return "PERF", "nfr"
            elif "non-functional" in title_lower or "nfr" in title_lower or "security" in title_lower:
                return "NFR", "nfr"
            elif "glossary" in title_lower:
                return "GLOSSARY", "general"
            else:
                clean_title = re.sub(r"^[0-9.\s]+", "", title).strip()
                slug = re.sub(r"[^\w]+", "_", clean_title.lower()).strip("_")[:15].upper()
                return slug or f"SEC-{index}", "general"

        # Full namespace prefix for Word XML paragraph style attribute.
        _WML = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        _PSTYLE_XPATH = f"{_WML}pPr/{_WML}pStyle"

        for child in doc.element.body:
            local_tag = child.tag.split("}")[-1]

            if local_tag == const.XML_PARAGRAPH_TAG:
                p_text = "".join(
                    node.text or ""
                    for node in child.iter()
                    if node.tag.endswith(const.XML_TEXT_TAG_SUFFIX)
                ).strip()
                if not p_text:
                    continue

                # Detect Word heading styles (Heading1, Heading2, etc.)
                is_heading = False
                style_node = child.find(_PSTYLE_XPATH)
                if style_node is not None:
                    val = style_node.attrib.get(f"{_WML}val", "")
                    if "Heading" in val:
                        is_heading = True

                if const.HEADING_REQUIREMENT_KEYWORD in p_text or is_heading:
                    sec_id_suffix, sec_type = derive_section_id(p_text, len(ast.sections))
                    sec_id = f"{module_folder}:{sec_id_suffix}"
                    current_section = SectionNode(
                        section_id=sec_id,
                        title=p_text,
                        type=sec_type,
                        module_folder=module_folder,
                        source_file=file_path.name,
                    )
                    ast.sections.append(current_section)
                else:
                    current_section.paragraphs.append(p_text)

            elif local_tag == const.XML_TABLE_TAG:
                table = DocxTable(child, doc)
                table_rows = []
                meta = current_section.metadata
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if len(cells) >= 2:
                        key = cells[0].lower().rstrip(":")
                        val = cells[1]

                        if const.KEY_DESCRIPTION in key:
                            meta["description"] = val
                        elif const.KEY_ACTORS in key:
                            meta["actors"] = [a.strip() for a in const.REGEX_TYPE_SPLIT.split(val) if a.strip()]
                        elif const.KEY_TRIGGER in key:
                            meta["trigger"] = val
                        elif const.KEY_PRIORITY in key:
                            meta["priority"] = val
                        elif any(k in key for k in const.KEY_PRECONDITIONS):
                            meta["pre_conditions"] = split_list_field(val)
                        elif const.KEY_BUSINESS_RULES in key:
                            meta["business_rules"] = split_list_field(val)
                        elif any(k in key for k in const.KEY_EXCEPTION_FLOW):
                            meta["exception_flows"] = split_numbered_steps(val)
                        elif const.KEY_MAIN_FLOW in key:
                            meta["main_flow"] = split_numbered_steps(val)
                        elif any(k in key for k in const.KEY_POSTCONDITIONS):
                            meta["post_conditions"] = split_list_field(val)
                        table_rows.append(cells)
                if table_rows:
                    current_section.tables.append(table_rows)

        return ast


# ---------------------------------------------------------------------------
# Test case parser
# ---------------------------------------------------------------------------

class TestCaseModuleParser:
    """Extracts all test cases from all tables in a .docx file."""

    @staticmethod
    def parse(file_path: Path, module_folder: str) -> List[TestCaseModel]:
        """
        Parse a Manual Test Cases .docx file and return all TestCaseModel objects.

        Iterates over every table in the document.  The first row of each table
        is treated as a header row; subsequent rows are data rows.  Tables without
        a "test name" column are skipped silently.

        Column discovery is dynamic (via ``find_col``) and supports alternative
        header spellings defined in ``core.constants`` (e.g. COL_TEST_NAME).

        The ``get_val`` helper is defined ONCE per table (not per row) to avoid
        redundant closure recreation inside the inner loop.

        Args:
            file_path:     Path to the .docx Test Cases file.
            module_folder: Name of the parent module folder (stored on each TC).

        Returns:
            Ordered list of TestCaseModel instances (may be empty if no valid
            test case tables are found or the file cannot be opened).
        """
        test_cases: List[TestCaseModel] = []
        try:
            doc = docx.Document(str(file_path))
        except Exception as exc:
            print(f"  [ERROR] Failed to read TC docx {file_path.name}: {exc}")
            return test_cases

        if not doc.tables:
            return test_cases

        for table_idx, table in enumerate(doc.tables):
            if not table.rows:
                continue

            header_row = table.rows[0]
            headers = [(cell.text or "").strip().lower() for cell in header_row.cells]

            def find_col(targets) -> Optional[int]:
                """Return the first column index whose header contains any target string."""
                targets_tuple = (targets,) if isinstance(targets, str) else targets
                for idx, h in enumerate(headers):
                    if any(t == h or (len(t) > 3 and t in h) for t in targets_tuple):
                        return idx
                return None

            col_map: Dict[str, Optional[int]] = {
                "test_name":       find_col(const.COL_TEST_NAME),
                "type":            find_col(["type", "category"]),
                "subject":         find_col(["subject", "module"]),
                "description":     find_col(const.COL_DESCRIPTION),
                "expected_result": find_col(const.COL_EXPECTED_RESULT),
                "execution_status":find_col(const.COL_EXECUTION_STATUS),
            }

            # Skip tables that have no recognisable "test name" column.
            if col_map["test_name"] is None:
                continue

            for row_idx, row in enumerate(table.rows[1:], start=1):
                cells = [(cell.text or "").strip() for cell in row.cells]

                # get_val is defined ONCE per table, not per row.
                # It closes over the current row's `cells` list via the outer scope.
                def get_val(key: str, _cells=cells) -> str:
                    """Safe column accessor -- returns empty string for missing columns."""
                    idx = col_map.get(key)
                    return _cells[idx] if idx is not None and idx < len(_cells) else ""

                tc_name_full = get_val("test_name")
                if not tc_name_full:
                    continue

                tc_id_match = const.REGEX_TC_ID.search(tc_name_full)
                tc_id       = tc_id_match.group(1).upper() if tc_id_match else f"TC-UNKN-{row_idx}"
                title_match = const.REGEX_TC_TITLE.search(tc_name_full)
                title       = title_match.group(1).strip() if title_match else tc_name_full

                type_str = get_val("type")
                types    = [t.strip() for t in const.REGEX_TYPE_SPLIT.split(type_str) if t.strip()]

                desc_text = get_val("description")
                steps     = split_numbered_steps(desc_text)
                if not steps and desc_text:
                    steps = [desc_text]

                tc = TestCaseModel(
                    tc_id=tc_id,
                    title=title,
                    module_folder=module_folder,
                    source_file=file_path.name,
                    source_table_index=table_idx,
                    source_row_index=row_idx,
                    type=types,
                    subject=get_val("subject"),
                    execution_status=get_val("execution_status"),
                    steps=steps,
                    expected_result=get_val("expected_result"),
                )
                test_cases.append(tc)

        return test_cases
